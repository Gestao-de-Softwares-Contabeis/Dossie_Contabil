import os
import tempfile
from io import BytesIO
from pathlib import Path
from typing import Optional

import fitz  # PyMuPDF
from docx import Document
from docx.shared import Inches

from app.utils.logger import logger
from config.settings import PDF_DPI, IMAGE_WIDTH_INCHES


class FileProcessor:
    """Responsável por conversão e inserção de arquivos no documento final."""

    # ------------------------------------------------------------------ #
    # Validação                                                             #
    # ------------------------------------------------------------------ #

    VALID_EXTENSIONS = {".docx", ".pdf", ".png", ".jpg", ".jpeg"}

    def validate_file(self, filename: str) -> bool:
        """Retorna True se a extensão do arquivo é suportada."""
        return Path(filename).suffix.lower() in self.VALID_EXTENSIONS

    # ------------------------------------------------------------------ #
    # PDF → imagens                                                         #
    # ------------------------------------------------------------------ #

    def pdf_to_images(self, pdf_path: str, max_pages: Optional[int] = None) -> list[bytes]:
        """
        Converte páginas de um PDF em bytes PNG.

        Args:
            pdf_path: Caminho do arquivo PDF.
            max_pages: Se informado, limita a quantidade de páginas convertidas.

        Returns:
            Lista de bytes PNG, uma por página.
        """
        images: list[bytes] = []
        try:
            doc = fitz.open(pdf_path)
            pages = doc[:max_pages] if max_pages else doc
            for page in pages:
                pix = page.get_pixmap(dpi=PDF_DPI)
                images.append(pix.tobytes("png"))
            logger.debug(f"PDF '{pdf_path}' convertido: {len(images)} página(s).")
        except Exception as exc:
            logger.error(f"Erro ao converter PDF '{pdf_path}': {exc}")
            raise
        return images

    def pdf_balanco_duas_paginas(self, pdf_path: str) -> tuple[bytes, bytes]:
        """
        Extrai as duas primeiras páginas do Balanço Patrimonial.

        Returns:
            Tupla (página1_bytes, página2_bytes).

        Raises:
            ValueError: Se o PDF tiver menos de 2 páginas.
        """
        images = self.pdf_to_images(pdf_path, max_pages=2)
        if len(images) < 2:
            raise ValueError(
                "O arquivo 'Balanço Patrimonial' deve ter pelo menos 2 páginas."
            )
        return images[0], images[1]

    # ------------------------------------------------------------------ #
    # Inserção de conteúdo em DOCX                                          #
    # ------------------------------------------------------------------ #

    def insert_pdf_at_placeholder(
        self, main_doc: Document, placeholder: str, pdf_path: str
    ) -> bool:
        """
        Substitui *placeholder* no documento por imagens geradas a partir do PDF.

        Returns:
            True se o placeholder foi encontrado e substituído.
        """
        images = self.pdf_to_images(pdf_path)
        for paragraph in main_doc.paragraphs:
            if placeholder in paragraph.text:
                paragraph.text = paragraph.text.replace(placeholder, "")
                for img_bytes in images:
                    run = paragraph.add_run()
                    run.add_picture(BytesIO(img_bytes), width=Inches(IMAGE_WIDTH_INCHES))
                logger.info(f"Placeholder '{placeholder}' substituído por {len(images)} imagem(ns).")
                return True
        logger.warning(f"Placeholder '{placeholder}' não encontrado no documento.")
        return False

    def insert_docx_at_placeholder(
        self, main_doc: Document, placeholder: str, insert_doc_path: str
    ) -> bool:
        """
        Substitui *placeholder* no documento inserindo o conteúdo de outro DOCX.

        Returns:
            True se o placeholder foi encontrado e substituído.
        """
        try:
            insert_doc = Document(insert_doc_path)
        except Exception as exc:
            logger.error(f"Erro ao abrir DOCX '{insert_doc_path}': {exc}")
            raise

        for paragraph in main_doc.paragraphs:
            if placeholder in paragraph.text:
                paragraph.text = paragraph.text.replace(placeholder, "")
                for element in reversed(list(insert_doc.element.body)):
                    paragraph._element.addnext(element)
                logger.info(f"Placeholder '{placeholder}' substituído com conteúdo de '{insert_doc_path}'.")
                return True

        logger.warning(f"Placeholder '{placeholder}' não encontrado no documento.")
        return False

    # ------------------------------------------------------------------ #
    # Utilitários                                                           #
    # ------------------------------------------------------------------ #

    @staticmethod
    def save_upload_to_temp(uploaded_file) -> str:
        """
        Salva um UploadedFile do Streamlit em arquivo temporário no disco.

        Returns:
            Caminho do arquivo temporário criado.
        """
        suffix = Path(uploaded_file.name).suffix
        with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp:
            tmp.write(uploaded_file.getvalue())
            return tmp.name

    @staticmethod
    def cleanup_temp_files(*paths: Optional[str]) -> None:
        """Remove arquivos temporários ignorando erros."""
        for path in paths:
            if path:
                try:
                    os.remove(path)
                except OSError:
                    pass