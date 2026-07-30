"""
Interface Streamlit do Gerador de Dossiês Contábeis.
Ponto de entrada para `streamlit run app/main.py`.
"""
import datetime
import json
import requests
import tempfile
import time
from pathlib import Path
from docx import Document

import sys
from pathlib import Path

raiz_projeto = Path(__file__).resolve().parent.parent
if str(raiz_projeto) not in sys.path:
    sys.path.append(str(raiz_projeto))

import streamlit as st

from core.automation import DossieAutomation
from utils.helpers import (
    format_cnpj,
    format_cpf,
    clean_numbers,
    build_periodo_em_data,
    build_periodo_anual,
    validate_cnpj,
    validate_cpf,
)
from config.constants import OUTPUT_FILENAME_TEMPLATE


def _coletar_erros_validacao(input_data: dict) -> list[str]:
    """Verifica todos os campos obrigatórios e retorna a lista de pendências/erros encontrados."""
    erros: list[str] = []

    if not input_data.get("nome_empresa", "").strip():
        erros.append("Nome Fantasia da Empresa")
    if not input_data.get("razao_social_empresa", "").strip():
        erros.append("Razão Social")

    cnpj_digitos = clean_numbers(input_data.get("cnpj_empresa", ""))
    if not cnpj_digitos:
        erros.append("CNPJ")
    elif len(cnpj_digitos) != 14:
        erros.append(f"CNPJ (deve ter 14 dígitos — tem {len(cnpj_digitos)})")
    elif not validate_cnpj(cnpj_digitos):
        erros.append("CNPJ (dígitos verificadores inválidos — confira o número)")

    if not input_data["uploads"].get("balanco_file"):
        erros.append("Balanço Patrimonial (PDF)")
    if not input_data["uploads"].get("demstr_result_file"):
        erros.append("Demonstração do Resultado — DRE (PDF)")

    socios_completos = 0
    for i, socio in enumerate(input_data.get("socios", []), start=1):
        nome = (socio.get("nome") or "").strip()
        cargo = (socio.get("cargo") or "").strip()
        cpf_digitos = clean_numbers(socio.get("cpf") or "")

        if not nome and not cargo and not cpf_digitos:
            continue  # linha em branco (ex: sócio recém-adicionado), ignora

        if not nome:
            erros.append(f"Sócio {i}: Nome completo")
        if not cargo:
            erros.append(f"Sócio {i}: Cargo")
        if not cpf_digitos:
            erros.append(f"Sócio {i}: CPF")
        elif len(cpf_digitos) != 11:
            erros.append(f"Sócio {i}: CPF (deve ter 11 dígitos — tem {len(cpf_digitos)})")
        elif not validate_cpf(cpf_digitos):
            erros.append(f"Sócio {i}: CPF (dígitos verificadores inválidos — confira o número)")
        elif nome and cargo:
            socios_completos += 1

    if socios_completos == 0:
        erros.append("Pelo menos um sócio completo (nome, CPF e cargo)")

    return erros

# Configuração da página
st.set_page_config(
    page_title="Gerador de Dossiês Contábeis",
    page_icon="📄",
    layout="wide",
)

st.title("📄 Gerador Automático de Dossiês Contábeis")
st.markdown("Preencha os campos e faça o upload dos arquivos para gerar o documento final.")

# Inicialização de estado 
if "socios" not in st.session_state:
    st.session_state.socios = [{"nome": "", "cpf": "", "cargo": ""}]

# Abas
tab1, tab2, tab3, tab4 = st.tabs(
    ["🏢 Empresa / Períodos", "👥 Sócios", "📎 Uploads", "⚠️ Pendências e Ressalvas"]
)

input_data: dict = {"uploads": {}}

# ── Aba 1 ─────────────────────────────────────────────────────────── #
with tab1:
    col1, col2 = st.columns(2)

    with col1:
        input_data["nome_empresa"] = st.text_input("Nome Fantasia da Empresa")
        input_data["razao_social_empresa"] = st.text_input("Razão Social")
        cnpj_raw = st.text_input("CNPJ (somente números)")
        # Mantemos apenas os dígitos no input_data: a formatação (XX.XXX.XXX/XXXX-XX)
        # é aplicada somente no momento da substituição no documento (automation.py),
        # não aqui — assim o CNPJ trafega "cru" por validação e envio ao n8n.
        input_data["cnpj_empresa"] = clean_numbers(cnpj_raw)
        if cnpj_raw:
            st.caption(f"Formatado: `{format_cnpj(cnpj_raw)}`")

    with col2:
        st.markdown("##### Período de Referência Contábil")
        today = datetime.date.today()
        data_inicio = st.date_input(
            "Data de Início",
            value=datetime.date(today.year, 1, 1),
            min_value=datetime.date(1980, 1, 1),
            max_value=today,
            key="data_inicio",
        )
        data_fim = st.date_input(
            "Data de Fim",
            value=today,
            min_value=datetime.date(1980, 1, 1),
            max_value=today,
            key="data_fim",
        )

        input_data["periodo_em_data"] = build_periodo_em_data(data_inicio, data_fim)
        input_data["periodo_anual"] = build_periodo_anual(data_inicio, data_fim)
        input_data["data_dem_encerradas"] = data_fim.strftime("%d/%m/%Y")

        st.markdown("---")
        st.caption(f"Período curto: `{input_data['periodo_em_data']}`")
        st.caption(f"Período longo: `{input_data['periodo_anual']}`")

# ── Aba 2 ─────────────────────────────────────────────────────────── #
with tab2:
    st.subheader("Dados dos Sócios / Administradores")

    for i, socio in enumerate(st.session_state.socios):
        with st.expander(f"Sócio {i + 1}", expanded=True):
            socio["nome"] = st.text_input(f"Nome completo", value=socio["nome"], key=f"nome_{i}")
            cpf_raw = st.text_input(f"CPF (somente números)", value=clean_numbers(socio["cpf"]), key=f"cpf_{i}")
            st.session_state.socios[i]["cpf"] = format_cpf(cpf_raw)
            if cpf_raw:
                st.caption(f"Formatado: `{st.session_state.socios[i]['cpf']}`")
            socio["cargo"] = st.text_input(f"Cargo", value=socio["cargo"], key=f"cargo_{i}")

            if st.button(f"🗑️ Remover Sócio {i + 1}", key=f"remove_{i}"):
                st.session_state.socios.pop(i)
                st.rerun()

    if st.button("➕ Adicionar Sócio"):
        st.session_state.socios.append({"nome": "", "cpf": "", "cargo": ""})
        st.rerun()

    input_data["socios"] = st.session_state.socios

# ── Aba 3 ─────────────────────────────────────────────────────────── #
with tab3:
    st.subheader("Balanço e DRE (PDF)")
    col5, col6 = st.columns(2)
    with col5:
        input_data["uploads"]["balanco_file"] = st.file_uploader(
            "Balanço Patrimonial (PDF — mínimo 2 páginas)", type=["pdf"]
        )
    with col6:
        input_data["uploads"]["demstr_result_file"] = st.file_uploader(
            "Demonstração do Resultado — DRE (PDF)", type=["pdf"]
        )

# ── Aba 4: Pendências (Substituindo o Telegram) ───────────────────── #
with tab4:
    st.subheader("Análise de Ressalvas e Pendências")
    st.markdown("Selecione se alguma pendência ou situação específica se aplica ao cenário atual da empresa.")
    
    OPCOES_PENDENCIAS = [
        "1. DISPONIBILIDADES – CAIXA", "2. DISPONIBILIDADES – CONTAS BANCÁRIAS", 
        "3. APLICAÇÕES FINANCEIRAS", "4. CONTAS A RECEBER (CLIENTES)", 
        "5. TRANSAÇÕES COM PARTES RELACIONADAS (SÓCIOS)", "6. ADIANTAMENTO A FORNECEDORES", 
        "7. ADIANTAMENTOS A EMPREGADOS", "8. RECUPERAÇÃO DE TRIBUTOS", 
        "9. SALDOS INICIAIS (BALANÇO DE ABERTURA)", "10. ESTOQUES", 
        "11. CONTROLE DE ESTOQUES EM PODER DE TERCEIROS", "12. ATIVO IMOBILIZADO", 
        "13. DEPRECIAÇÃO E VIDA ÚTIL (CPC 27)", "14. FORNECEDORES (PASSIVO CIRCULANTE)", 
        "15. OBRIGAÇÕES TRIBUTÁRIAS", "16. CONCILIAÇÃO DE FOLHA DE PAGAMENTO", 
        "17. CONFUSÃO PATRIMONIAL (PRINCÍPIO DA ENTIDADE)", "18. PASSIVO A DESCOBERTO / OMISSÃO DE RECEITA", 
        "19. DISPÊNDIOS SEM COMPROVAÇÃO FISCAL", "20. PASSIVOS ONEROSOS (EMPRÉSTIMOS)", 
        "21. CONSÓRCIOS", "22. DESPESAS ANTECIPADAS (SEGUROS)", 
        "23. PARCELAMENTOS TRIBUTÁRIOS", "24. ARRENDAMENTO MERCANTIL E ALUGUÉIS (CPC 06)", 
        "25. OBRIGAÇÕES COM CARTÕES DE CRÉDITO", "Nenhuma das Opções"
    ]
    
    input_data["pendencias"] = st.multiselect(
        "Selecione as pendências ou ressalvas aplicáveis:", 
        options=OPCOES_PENDENCIAS,
        default=["Nenhuma das Opções"]
    )

# Botão de geração
st.divider()
if st.button("✅ GERAR DOSSIÊ CONTÁBIL", type="primary", use_container_width=True):
    erros_validacao = _coletar_erros_validacao(input_data)

    if erros_validacao:
        lista = "\n".join(f"- {erro}" for erro in erros_validacao)
        st.warning(f"⚠️ Corrija os campos abaixo antes de gerar o dossiê:\n\n{lista}")
    else:
        automation = DossieAutomation()

        # Padrão assíncrono: o POST inicial só cria o job e recebe um job_id de volta
        # rapidamente (evita o timeout do gateway/proxy na frente do n8n). O processamento
        # pesado (extração de PDF, chamadas de IA e o loop de reavaliação de qualidade)
        # roda em background no n8n; consultamos o resultado por polling em N8N_STATUS_URL.
        N8N_WEBHOOK_URL = "https://genai.up4me.io/webhook/receber-dados"
        N8N_STATUS_URL = "https://genai.up4me.io/webhook/status-dossie"
        POLL_INTERVAL_SECONDS = 5
        POLL_TIMEOUT_SECONDS = 600  # margem para o loop de retry de qualidade do workflow

        status_placeholder = st.empty()
        status_placeholder.info("📤 Enviando dados ao n8n...")

        payload = {
            "nome_empresa": input_data.get("nome_empresa", ""),
            "razao_social_empresa": input_data.get("razao_social_empresa", ""),
            "cnpj_empresa": input_data.get("cnpj_empresa", ""),
            "periodo_em_data": input_data.get("periodo_em_data", ""),
            "periodo_anual": input_data.get("periodo_anual", ""),
            "data_dem_encerradas": input_data.get("data_dem_encerradas", ""),
            # "pendencias" é lista de strings: enviada como campos multipart repetidos,
            # o n8n reconstrói isso como array automaticamente — não precisa de json.dumps.
            "pendencias": input_data.get("pendencias", []),
            # "socios" é lista de DICTS: junto de "files" em multipart, o requests não
            # converte isso em JSON sozinho, ele gera repr Python (aspas simples), que
            # quebra o JSON.parse() do lado do n8n. Por isso serializamos manualmente
            # (o workflow deve dar JSON.parse() em body.socios).
            "socios": json.dumps(input_data.get("socios", []))
        }

        # Extraímos os objetos de arquivo do Streamlit
        balanco_file = input_data["uploads"]["balanco_file"]
        dre_file = input_data["uploads"]["demstr_result_file"]

        # (nome_do_arquivo, conteudo_em_bytes, tipo_mime)
        files = {
            "balanco": (balanco_file.name, balanco_file.getvalue(), balanco_file.type),
            "dre": (dre_file.name, dre_file.getvalue(), dre_file.type)
        }

        job_id = None
        dados_gerados = None
        erro_fatal = None

        # 1. POST inicial: deve ser rápido — só cria o job e devolve {"job_id": ...}.
        # Timeout curto de propósito, pois isso NÃO deve esperar o processamento pesado.
        try:
            response = requests.post(N8N_WEBHOOK_URL, data=payload, files=files, timeout=30)

            if not response.ok:
                erro_fatal = f"❌ Erro do n8n: {response.status_code} - {response.text}"
            elif not response.text.strip():
                erro_fatal = (
                    "❌ O n8n respondeu com sucesso (200), mas sem corpo. Verifique se o "
                    "node 'Respond to Webhook' logo após criar o job está retornando "
                    "{'job_id': ...} — ele deve ficar ANTES do processamento pesado, não no final."
                )
            else:
                try:
                    ack = response.json()
                except ValueError:
                    erro_fatal = (
                        "❌ A resposta inicial do n8n não é um JSON válido. "
                        f"Corpo recebido (primeiros 500 caracteres): {response.text[:500]!r}"
                    )
                else:
                    if isinstance(ack, list):
                        ack = ack[0] if ack else {}
                    job_id = ack.get("job_id")
                    if not job_id:
                        erro_fatal = f"❌ O n8n não retornou 'job_id' na resposta inicial: {ack!r}"

        except requests.exceptions.Timeout:
            erro_fatal = (
                "❌ O n8n demorou mais de 30s só para confirmar o recebimento do pedido "
                "(criar o job). Isso não deveria envolver processamento pesado — confira "
                "se o node 'Respond to Webhook' inicial está mesmo antes das chamadas de IA."
            )
        except requests.exceptions.ConnectionError as e:
            erro_fatal = (
                "❌ O n8n encerrou a conexão sem responder ao POST inicial (Connection aborted). "
                f"Detalhes técnicos: {e}"
            )
        except Exception as e:
            erro_fatal = f"Erro na comunicação com n8n: {e}"

        # 2. Polling: aguarda o job concluir em background, sem manter uma única conexão
        # HTTP aberta por minutos (evitando o timeout do proxy/gateway na frente do n8n).
        if job_id and not erro_fatal:
            elapsed = 0
            while elapsed < POLL_TIMEOUT_SECONDS:
                status_placeholder.info(
                    f"⏳ Processando no n8n... ({elapsed}s / até {POLL_TIMEOUT_SECONDS}s) — job {job_id}"
                )
                try:
                    status_resp = requests.get(
                        N8N_STATUS_URL, params={"job_id": job_id}, timeout=15
                    )
                    status_resp.raise_for_status()
                    status_json = status_resp.json()
                    if isinstance(status_json, list):
                        status_json = status_json[0] if status_json else {}
                except Exception as e:
                    erro_fatal = f"❌ Falha ao consultar status do job {job_id}: {e}"
                    break

                status = status_json.get("status")
                if status == "done":
                    dados_gerados = status_json.get("resultado", {}) or {}
                    break
                elif status == "error":
                    erro_fatal = (
                        f"❌ O n8n reportou erro no processamento do job {job_id}: "
                        f"{status_json.get('erro', 'sem detalhes')}"
                    )
                    break

                time.sleep(POLL_INTERVAL_SECONDS)
                elapsed += POLL_INTERVAL_SECONDS
            else:
                erro_fatal = (
                    f"❌ Tempo limite de espera ({POLL_TIMEOUT_SECONDS}s) excedido "
                    f"aguardando o n8n concluir o job {job_id}."
                )

        status_placeholder.empty()

        if erro_fatal:
            st.error(erro_fatal)
        elif dados_gerados is not None:
            notas_raw = dados_gerados.get("explic_demonstr", "")
            try:
                # "explic_demonstr" vem como uma string JSON aninhada:
                # '{"explic_demonstr": "texto..."}' — desembrulha se for o caso
                notas_texto = json.loads(notas_raw).get("explic_demonstr", notas_raw)
            except (TypeError, ValueError):
                notas_texto = notas_raw or "Erro ao gerar notas"

            carta_texto = dados_gerados.get("carta_responsb", "Erro ao gerar carta")

            st.success("✨ Textos gerados pelo n8n com sucesso! Montando dossiê...")

            # Converte os textos retornados em arquivos DOCX temporários para o Automation
            doc_notas = Document()
            doc_notas.add_paragraph(notas_texto)
            path_notas = Path(tempfile.gettempdir()) / "temp_notas.docx"
            doc_notas.save(path_notas)

            doc_carta = Document()
            doc_carta.add_paragraph(carta_texto)
            path_carta = Path(tempfile.gettempdir()) / "temp_carta.docx"
            doc_carta.save(path_carta)

            # Injeta os caminhos no input_data para simular o upload que foi comentado
            input_data["uploads"]["explic_demonstr_file"] = str(path_notas)
            input_data["uploads"]["carta_responsb_file"] = str(path_carta)

            # Gera o arquivo final
            file_bytes, error = automation.generate(input_data)

            if error:
                st.error(f"❌ Falha na montagem final: {error}")
            else:
                st.success("✅ Dossiê completo montado!")
                filename = OUTPUT_FILENAME_TEMPLATE.format(nome_empresa=input_data["nome_empresa"])
                st.download_button(
                    label="⬇️ Baixar Dossiê Completo (.docx)",
                    data=file_bytes,
                    file_name=filename,
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )